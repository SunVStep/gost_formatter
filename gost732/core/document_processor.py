from docx import Document
from pathlib import Path

class DocumentProcessor:

    def _copy_content(self, source_doc, target_doc):

        for par in source_doc.paragraphs:
            if not par.text.strip():
                continue
            new_p = target_doc.add_paragraph()
            target_style = 'ГОСТ_Текст' if 'ГОСТ_Текст' in target_doc.styles else 'Normal'
            new_p.style = target_style

            for run in par.runs:
                new_run = new_p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

    def process(self, input_path: Path, template_path: Path, output_path: Path):

        source_doc = Document(input_path)
        target_doc = Document(template_path)

        self._copy_content(source_doc, target_doc)

        target_doc.save(output_path)