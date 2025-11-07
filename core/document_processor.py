from docx import Document
from pathlib import Path

project_path = Path(__file__).resolve().parent.parent
reports_dir = project_path / 'reports'
templates_dir = project_path / 'template'
output_dir = project_path / 'output'

dirty_path = reports_dir / 'test.docx'
template_path = templates_dir / 'Main_Template.docx'

if not dirty_path.exists():
    raise FileNotFoundError(f"Исходный файл не найден: {dirty_path}")
if not template_path.exists():
    raise FileNotFoundError(f"Файл шаблона не найден: {template_path}")

dirty_doc = Document(dirty_path)
final_document = Document(template_path)

for paragraphs in dirty_doc.paragraphs:
    if paragraphs.text.strip():
        added = final_document.add_paragraph(paragraphs.text)
        if 'ГОСТ_Текст' in final_document.styles:
            added.style = 'ГОСТ_Текст'
        else:
            added.style = 'Normal'

output_path = project_path / 'output' / 'Final_report.docx'
final_document.save(output_path)
print(f"✅ Успешно! Результат сохранён: {output_path}")
